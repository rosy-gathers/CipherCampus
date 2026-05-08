from pathlib import Path
import re
from docx import Document


SRC = Path("CSE447_Report_Completed.md")
OUT = Path("CSE447_Report_Completed.docx")


def add_markdown_paragraph(doc: Document, line: str) -> None:
    stripped = line.strip()

    if not stripped:
        doc.add_paragraph("")
        return

    if stripped.startswith("# "):
        doc.add_heading(stripped[2:].strip(), level=1)
        return
    if stripped.startswith("## "):
        doc.add_heading(stripped[3:].strip(), level=2)
        return
    if stripped.startswith("### "):
        doc.add_heading(stripped[4:].strip(), level=3)
        return

    if stripped in {"---", "***"}:
        return

    if re.match(r"^\d+\.\s", stripped):
        doc.add_paragraph(stripped, style="List Number")
        return

    if stripped.startswith("- "):
        doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        return

    if stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
        p = doc.add_paragraph()
        p.add_run(stripped[2:-2]).bold = True
        return

    # Light inline markdown cleanup
    cleaned = stripped.replace("**", "").replace("`", "")
    doc.add_paragraph(cleaned)


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(f"Source markdown not found: {SRC}")

    lines = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()

    for line in lines:
        add_markdown_paragraph(doc, line)

    doc.save(OUT)
    print(f"Created: {OUT.resolve()}")


if __name__ == "__main__":
    main()
